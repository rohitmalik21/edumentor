"""
EduMentor AI - Document Parser
Extracts text from PDF, DOCX, and TXT files.
Handles large files by extracting meaningful content (skipping front/back matter).
"""

import os
import re
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text content from a file based on its extension.

    Args:
        file_path: Path to the uploaded file.

    Returns:
        Extracted text content as a string.

    Raises:
        ValueError: If file format is not supported.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    elif ext in (".txt", ".md"):
        return _extract_from_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: PDF, DOCX, TXT")


def _extract_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    For large PDFs (books), skips front/back matter and extracts body content.
    """
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)

    # For large PDFs (likely books), skip front matter
    if total_pages > 20:
        # Skip first 5-10 pages (cover, copyright, TOC, preface)
        # and last few pages (index, ads)
        skip_start = min(8, total_pages // 10)
        skip_end = min(5, total_pages // 20)
        pages_to_read = reader.pages[skip_start:total_pages - skip_end]

        # For very large books, sample pages evenly to get broad coverage
        if len(pages_to_read) > 50:
            # Take content from beginning, middle, and spread across chapters
            step = max(1, len(pages_to_read) // 30)
            pages_to_read = pages_to_read[::step]
    else:
        pages_to_read = reader.pages

    text_parts = []
    for page in pages_to_read:
        page_text = page.extract_text()
        if page_text:
            # Clean up the text
            cleaned = _clean_extracted_text(page_text)
            if cleaned and len(cleaned) > 50:  # Skip near-empty pages
                text_parts.append(cleaned)

    full_text = "\n\n".join(text_parts)

    # Remove common noise from book PDFs
    full_text = _remove_book_noise(full_text)

    return full_text


def _extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file, including tables."""
    doc = Document(file_path)

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_rows.append(row_text)
        if table_rows:
            text_parts.append("\n".join(table_rows))

    return "\n\n".join(text_parts)


def _extract_from_text(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _clean_extracted_text(text: str) -> str:
    """Clean up text extracted from PDF pages."""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove page numbers standalone
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    # Remove URLs
    text = re.sub(r'https?://\S+', '', text)
    # Remove email addresses
    text = re.sub(r'\S+@\S+\.\S+', '', text)
    return text.strip()


def _remove_book_noise(text: str) -> str:
    """Remove common publisher/book noise from extracted text."""
    # Patterns to remove
    noise_patterns = [
        r'(?i)all rights reserved.*?(?=\n|$)',
        r'(?i)copyright\s*©.*?(?=\n|$)',
        r'(?i)published by.*?(?=\n|$)',
        r'(?i)isbn[\s:-]*[\d-]+',
        r'(?i)packt publishing.*?(?=\n|$)',
        r'(?i)www\.packt\.com.*?(?=\n|$)',
        r'(?i)printed in.*?(?=\n|$)',
        r'(?i)first published.*?(?=\n|$)',
        r'(?i)production reference.*?(?=\n|$)',
        r'(?i)disclaimer.*?(?=\n|$)',
        r'(?i)table of contents',
        r'(?i)about the author.*?(?=\n|$)',
        r'(?i)about the reviewer.*?(?=\n|$)',
        r'(?i)preface\s*$',
    ]

    for pattern in noise_patterns:
        text = re.sub(pattern, '', text)

    # Remove lines that are just numbers (page numbers)
    text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)

    # Remove multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Split text into overlapping chunks for RAG processing.

    Args:
        text: Full text to chunk.
        chunk_size: Maximum characters per chunk.
        overlap: Number of overlapping characters between chunks.

    Returns:
        List of text chunks.
    """
    if not text:
        return []

    chunks = []
    sentences = text.replace("\n", " ").split(". ")
    current_chunk = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        if len(current_chunk) + len(sentence) + 2 <= chunk_size:
            current_chunk += sentence + ". "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # Start new chunk with overlap from previous
            if overlap > 0 and current_chunk:
                overlap_text = current_chunk[-overlap:]
                current_chunk = overlap_text + sentence + ". "
            else:
                current_chunk = sentence + ". "

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks
